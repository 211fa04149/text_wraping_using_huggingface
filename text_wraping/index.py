import os
import requests
from dotenv import load_dotenv
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from datetime import datetime
from textwrap import wrap
from docx import Document
from docx.shared import Pt


# Load env
load_dotenv()

API_TOKEN = os.getenv("HF_API_TOKEN")
if not API_TOKEN:
    print("❌ API token not found")
    exit()

API_URL = "https://router.huggingface.co/v1/chat/completions"

HEADERS = {
    "Authorization": f"Bearer {API_TOKEN}",
    "Content-Type": "application/json"
}

# ---------------- AI FUNCTION ----------------
def generate_assignment(topic):
    prompt = f"""
Explain the topic "{topic}" clearly.

Then generate:
1. Assignment Questions with difficulty levels:
   - Easy (2 questions)
   - Medium (2 questions)
   - Hard (2 questions)
   questions must be descriptive and theory-based

2. MCQs (5 questions) with 4 options each.
   Do NOT give answers.

Use clear headings:
Explanation
Assignment Questions
MCQs
"""

    payload = {
        "model": "meta-llama/Meta-Llama-3-8B-Instruct",
        "messages": [
            {"role": "system", "content": "You are a professional teacher preparing student assignments."},
            {"role": "user", "content": prompt}
        ],
        "max_tokens": 900,
        "temperature": 0.6
    }

    response = requests.post(API_URL, headers=HEADERS, json=payload)

    if response.status_code != 200:
        return None, f"❌ HTTP Error: {response.text}"

    data = response.json()
    return data["choices"][0]["message"]["content"], None

# ---------------- PDF FUNCTION ----------------
def save_to_pdf(topic, content):
    filename = f"{topic.replace(' ', '_')}_assignment.pdf"
    c = canvas.Canvas(filename, pagesize=A4)
    width, height = A4

    x_margin = 40
    y = height - 50

    c.setFont("Times-Bold", 16)
    c.drawString(x_margin, y, f"Assignment Topic: {topic}")
    y -= 30

    c.setFont("Times-Roman", 10)
    c.drawString(x_margin, y, f"Generated on: {datetime.now().strftime('%d-%m-%Y %H:%M')}")
    y -= 30

    c.setFont("Times-Roman", 11)

    for line in content.split("\n"):
        if y < 50:
            c.showPage()
            c.setFont("Times-Roman", 11)
            y = height - 50
        c.drawString(x_margin, y, line)
        y -= 15

    c.save()
    return filename




# ---------------- WORD SAVE ----------------
def save_to_word(topic, content):
    os.makedirs("outputs", exist_ok = True)
    filename = f"outputs/{topic.replace(' ', '_')}_assignment.docx"
    doc = Document()

    title = doc.add_heading(f"Assignment Topic: {topic}", level=1)
    title.alignment = 1

    

    doc.add_paragraph(
        f"Generated on: {datetime.now().strftime('%d-%m-%Y %H:%M')}"
    )
    doc.add_paragraph("")
    
    for line in content.split("\n"):
        p = doc.add_paragraph(line)
        for run in p.runs:
            run.font.size = Pt(11)

    doc.save(filename)
    return filename

# ---------------- MAIN LOOP ----------------
print("📘 Enter a topic to generate Assignment (type 'exit' to quit)")

while True:
    topic = input("\nUser: ").strip()

    if topic.lower() == "exit":
        print("👋 Exiting program.")
        break

    if not topic:
        print("⚠️ Please enter a valid topic.")
        continue

    result, error = generate_assignment(topic)

    if error:
        print(error)
        continue

    print("\n📄 Assignment Generated:\n")
    print(result)

    pdf_file = save_to_pdf(topic, result)
    word_file = save_to_word(topic, result)
    print(f"\n✅ Saved as PDF: {pdf_file}")

    
    #word_file = save_to_word(topic, result)
    print(f"✅ Word saved as: {word_file}")